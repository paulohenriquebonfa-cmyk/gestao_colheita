from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
EBOOK_PATH = DOCS_DIR / "Ebook_Guia_de_Uso_Gestao_de_Colheita.docx"
QUICK_PATH = DOCS_DIR / "Guia_Rapido_Operador_Gestao_de_Colheita.docx"

BLUE = "173F5F"
MID_BLUE = "2E74B5"
DARK = "1F2937"
MUTED = "5B6573"
LIGHT_FILL = "EEF5FB"
SOFT_FILL = "F6F9FC"
LINE = "D7E2EC"


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in (
        ("Heading 1", 16, MID_BLUE),
        ("Heading 2", 13, MID_BLUE),
        ("Heading 3", 12, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run("Gestao de Colheita | Guia do usuario")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc, title, subtitle, audience):
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    banner.columns[0].width = Inches(6.5)
    style_table(banner)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, BLUE)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=10, after=8, line=1.0)
    run = p.add_run("GESTAO DE COLHEITA")
    set_run_font(run, size=12, bold=True, color="FFFFFF")

    p2 = cell.add_paragraph()
    set_spacing(p2, before=0, after=4, line=1.0)
    run2 = p2.add_run(title)
    set_run_font(run2, size=24, bold=True, color="FFFFFF")

    p3 = cell.add_paragraph()
    set_spacing(p3, before=0, after=10, line=1.15)
    run3 = p3.add_run(subtitle)
    set_run_font(run3, size=11, color="EAF3FA")

    p4 = cell.add_paragraph()
    set_spacing(p4, before=6, after=6, line=1.0)
    run4 = p4.add_run(audience)
    set_run_font(run4, size=10, bold=True, color="FFFFFF")

    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_spacing(p, before=12 if level == 1 else 8, after=5 if level == 1 else 4, line=1.1)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=MID_BLUE if level < 3 else BLUE)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_spacing(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)


def add_note(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    style_table(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, LIGHT_FILL)

    p = cell.paragraphs[0]
    set_spacing(p, before=0, after=3, line=1.0)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True, color=BLUE)

    for line in lines:
        p = cell.add_paragraph()
        set_spacing(p, before=0, after=2, line=1.1)
        run = p.add_run(line)
        set_run_font(run, size=10, color=DARK)
    doc.add_paragraph()


def add_screen_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    style_table(table)

    for i, value in enumerate(("Tela", "Para que serve")):
        cell = table.rows[0].cells[i]
        shade_cell(cell, SOFT_FILL)
        p = cell.paragraphs[0]
        set_spacing(p, before=0, after=0)
        run = p.add_run(value)
        set_run_font(run, size=11, bold=True, color=BLUE)

    rows = [
        ("Dashboard", "Mostra os números principais da safra e os gráficos."),
        ("Nova Carga", "Registra cada carga da colheita."),
        ("Historico", "Consulta, filtra, edita ou apaga cargas."),
        ("Cadastros", "Organiza produtores, propriedades, talhões, variedades, armazéns e caminhões."),
        ("Analises", "Compara produtividade e desempenho."),
        ("Frete", "Controla tarifas, diesel, vale, conferência e recibo."),
        ("Armazenagem e Vendas", "Mostra saldo, registra saídas e vendas."),
        ("Assistente", "Ajuda com backup, dados de teste e configuração.")
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((left, right)):
            p = cells[idx].paragraphs[0]
            set_spacing(p, before=0, after=0)
            run = p.add_run(value)
            set_run_font(run, size=10, color=DARK)
    doc.add_paragraph()


def build_ebook():
    doc = Document()
    configure_document(doc)

    add_cover(
        doc,
        "Guia de Uso do Sistema",
        "Material simples, direto e fácil de entender para qualquer pessoa usar o sistema no dia a dia.",
        "Indicado para proprietários, operadores e pessoas do escritório."
    )

    add_heading(doc, "1. O que este sistema faz")
    add_paragraph(doc, "O sistema ajuda a organizar a colheita, o frete, o estoque e as vendas. Ele também separa tudo por safra, para que os dados de um ano não se misturem com os de outro.")
    add_bullets(doc, [
        "Controla os cadastros principais da operação.",
        "Registra cargas da colheita.",
        "Mostra números e gráficos no dashboard.",
        "Controla frete, diesel, vale e recibo.",
        "Controla estoque, armazenagem e vendas.",
        "Permite consultar safras anteriores."
    ])

    add_note(doc, "Comece por aqui", [
        "Se você está entrando pela primeira vez, leia nesta ordem: Cadastros, Safras e Nova Carga.",
        "Depois siga para Frete, Armazenagem e Análises."
    ])

    add_heading(doc, "2. Primeiro acesso")
    add_numbers(doc, [
        "Abra o sistema e faça login.",
        "Olhe a safra mostrada no topo da tela.",
        "Se a safra estiver errada, troque no seletor antes de lançar qualquer informação.",
        "Se for a primeira vez, cadastre os dados básicos da operação."
    ])

    add_heading(doc, "3. Entendendo as telas principais")
    add_screen_table(doc)

    add_heading(doc, "4. O que cadastrar primeiro")
    add_paragraph(doc, "Antes de lançar cargas, deixe os cadastros organizados.")
    add_bullets(doc, [
        "Propriedades: origem da produção.",
        "Produtores: quem produz ou recebe a produção.",
        "Talhões: áreas da fazenda.",
        "Variedades: soja, milho e outros materiais.",
        "Armazéns: locais para onde a produção vai.",
        "Caminhões: veículos usados no transporte."
    ])

    add_note(doc, "Dica prática", [
        "Use nomes simples e padronizados.",
        "Exemplo: Talhao Norte, Silo Sede, Cooperativa, Caminhao BXC9D09."
    ])

    add_heading(doc, "5. Como trabalhar com safras")
    add_paragraph(doc, "A safra é o centro do sistema. Quase tudo fica ligado a ela: carga, frete, estoque, venda e relatórios.")
    add_numbers(doc, [
        "Abra a tela de Frete.",
        "Na parte Cadastro de Safra, informe nome, cultura, ano e período.",
        "Se for a safra que você vai usar agora, marque como ativa.",
        "Se a safra já começar com saldo em armazém, informe o saldo inicial.",
        "Salve a safra."
    ])
    add_bullets(doc, [
        "A safra ativa aparece no topo do sistema.",
        "Ao trocar a safra, as telas passam a mostrar só os dados daquela safra.",
        "Safras antigas continuam guardadas para consulta."
    ])

    add_heading(doc, "6. Como lançar uma nova carga")
    add_numbers(doc, [
        "Escolha a data da carga.",
        "Selecione caminhão, propriedade, talhão, produtor, variedade e armazém.",
        "Informe peso bruto e peso líquido.",
        "Confira o frete quando a rota tiver tarifa cadastrada.",
        "Salve a carga."
    ])
    add_note(doc, "Muito importante", [
        "Peso líquido é usado para produção, análises e estoque.",
        "Peso bruto é usado para o cálculo do frete."
    ])

    add_heading(doc, "7. Como usar o Dashboard")
    add_bullets(doc, [
        "Veja total líquido colhido.",
        "Veja total em sacas.",
        "Acompanhe produtividade geral.",
        "Observe os gráficos por dia, talhão e produtor.",
        "Lembre que os números sempre dependem da safra em exibição."
    ])

    add_heading(doc, "8. Como usar o Histórico")
    add_paragraph(doc, "No Histórico você consulta tudo o que já foi lançado.")
    add_bullets(doc, [
        "Pode filtrar por data, produtor, propriedade, talhão, variedade e armazém.",
        "Pode editar uma carga se algo foi digitado errado.",
        "Pode apagar uma carga, mas isso muda estoque e relatórios da safra."
    ])

    add_heading(doc, "9. Como usar as Análises")
    add_bullets(doc, [
        "Compare produtividade por talhão.",
        "Veja o desempenho por variedade.",
        "Use para entender quais áreas estão produzindo melhor.",
        "Sempre confira a safra selecionada antes de comparar."
    ])

    add_heading(doc, "10. Como controlar o Frete")
    add_paragraph(doc, "A tela de Frete serve para controlar tarifas, abastecimentos, vales e documentos.")
    add_numbers(doc, [
        "Selecione a safra e o caminhão.",
        "Cadastre a tarifa da rota por propriedade e armazém.",
        "Registre diesel quando o caminhoneiro abastecer na propriedade.",
        "Registre vale quando houver adiantamento em dinheiro.",
        "No final, gere o relatório de conferência ou o recibo."
    ])
    add_bullets(doc, [
        "O cálculo do frete usa o peso bruto.",
        "Diesel e vale entram como desconto do frete.",
        "O relatório de conferência é diferente do recibo."
    ])

    add_heading(doc, "11. Como controlar Armazenagem e Vendas")
    add_numbers(doc, [
        "Abra a tela Armazenagem e Vendas.",
        "Confira o saldo por armazém da safra atual.",
        "Registre a venda informando data, produtor, armazém e valor por saca.",
        "Se precisar corrigir, use ajuste manual com cuidado.",
        "Se uma venda for cancelada, use o cancelamento para o sistema estornar corretamente."
    ])

    add_heading(doc, "12. Como usar o Assistente")
    add_bullets(doc, [
        "Exportar Backup: salva seus dados em arquivo.",
        "Importar Backup: restaura dados salvos.",
        "Salvar snapshot semanal: guarda uma cópia rápida.",
        "Carregar dados de teste: cria um ambiente para testar sem digitar tudo.",
        "Excluir dados de teste: apaga apenas os dados de teste."
    ])

    add_heading(doc, "13. Sincronização entre aparelhos")
    add_numbers(doc, [
        "Faça login com a conta correta.",
        "Use o sistema normalmente mesmo se a internet oscilar.",
        "Quando estiver online, clique em Sincronizar.",
        "Se aparecer erro, leia a mensagem e tente novamente depois."
    ])
    add_note(doc, "Regra simples", [
        "Se o dado salvou no aparelho, ele pode ser sincronizado depois.",
        "Mas sempre confira a safra ativa antes de lançar uma nova informação."
    ])

    add_heading(doc, "14. Rotina recomendada")
    add_numbers(doc, [
        "Conferir a safra ativa.",
        "Lançar as cargas do dia.",
        "Registrar diesel e vales, quando houver.",
        "Conferir dashboard e histórico.",
        "Verificar armazenagem e vendas.",
        "Sincronizar no fim do dia."
    ])

    add_heading(doc, "15. Dúvidas comuns")
    add_bullets(doc, [
        "Os números mudaram? Veja se a safra selecionada está correta.",
        "O frete parece diferente? Lembre que ele usa o peso bruto.",
        "O estoque não bate? Revise cargas, vendas e ajustes da safra.",
        "Vai começar outro ano? Crie uma nova safra e deixe ela ativa."
    ])

    add_heading(doc, "16. Fechamento")
    add_paragraph(doc, "O sistema foi pensado para facilitar o trabalho no campo e no escritório. Se você mantiver os cadastros organizados, escolher a safra certa e lançar as cargas com atenção, o restante do controle fica muito mais simples.")

    doc.save(EBOOK_PATH)


def build_quick_guide():
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    add_cover(
        doc,
        "Guia Rápido do Operador",
        "Resumo de uso diário para lançar informações sem erro.",
        "Leitura rápida para o trabalho do dia a dia."
    )

    add_heading(doc, "Antes de começar", level=1)
    add_bullets(doc, [
        "Confira a safra no topo da tela.",
        "Veja se está na safra certa antes de lançar qualquer coisa.",
        "Se faltar algum cadastro, peça para cadastrar antes."
    ])

    add_heading(doc, "Ordem mais segura de trabalho", level=1)
    add_numbers(doc, [
        "Abrir Nova Carga.",
        "Lançar as cargas do dia.",
        "Conferir Histórico.",
        "Registrar diesel e vale no Frete, se houver.",
        "Conferir Armazenagem e Vendas.",
        "Sincronizar no fim do dia."
    ])

    add_heading(doc, "Na hora de lançar uma carga", level=1)
    add_bullets(doc, [
        "Preencha data, caminhão, propriedade, talhão, produtor, variedade e armazém.",
        "Digite peso bruto e peso líquido com atenção.",
        "Salve a carga só depois de conferir."
    ])

    add_note(doc, "Lembrete", [
        "Peso líquido = produção e estoque.",
        "Peso bruto = frete."
    ])

    add_heading(doc, "Na hora de usar o Frete", level=1)
    add_bullets(doc, [
        "Selecione a safra e o caminhão.",
        "Registre diesel sempre que houver abastecimento.",
        "Registre vale sempre que houver adiantamento.",
        "Use recibo só quando for pagamento."
    ])

    add_heading(doc, "Se der dúvida", level=1)
    add_bullets(doc, [
        "Veja se a safra certa está selecionada.",
        "Confirme os cadastros da carga.",
        "Revise o Histórico antes de corrigir algo.",
        "Se precisar, use o ebook completo."
    ])

    doc.save(QUICK_PATH)


def main():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    build_ebook()
    build_quick_guide()
    print(EBOOK_PATH)
    print(QUICK_PATH)


if __name__ == "__main__":
    main()
